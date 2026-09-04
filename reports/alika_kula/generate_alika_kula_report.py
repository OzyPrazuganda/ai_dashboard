from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd


OUTPUT = "c:\\ai_dashboard\\reports\\alika_kula\\ALIKA_KULA_Journal_Research_Report.docx"
FIGURE_DIR = "c:\\ai_dashboard\\reports\\alika_kula\\figures"


def build_figures():
    import os
    os.makedirs(FIGURE_DIR, exist_ok=True)
    files = {
        "May": ["akulaku_superset_presto_hanmm_ozy.prazuganda_1788252974 (1-5 may).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788253288 (6-10 may).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788318709 (11-15 may).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788318958 (16-20 may).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788319766 (21-25 may).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788321413 (26-31 may).csv"],
        "June": ["akulaku_superset_presto_hanmm_ozy.prazuganda_1788146480 (1-5 june).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788146837 (6-10 june).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788148372 (11-15 june).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788149085 (16-20 june).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788149522 (21-25 june).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788150625 (26-30 june).csv"],
        "July": ["akulaku_superset_presto_hanmm_ozy.prazuganda_1788244584 (1-5 july).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788244938 (6-10 july).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788247515 (11-15 july).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788248125 (16-20 july).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788252178 (21-25 july).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788252384 (26-31 july).csv"],
        "August": ["akulaku_superset_presto_hanmm_ozy.prazuganda_1788145507 (1-5 aug).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788145209 (6-10 aug).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788144961 (11-14 aug).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1787219325 (13-15 aug).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1787025549 (16-18 aug).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1787219503 (19-20 aug).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788144629 (21-26 aug).csv", "akulaku_superset_presto_hanmm_ozy.prazuganda_1788141789 (27-31 aug).csv"],
    }
    raw = {month: pd.concat([pd.read_csv("c:\\ai_dashboard\\dataset_wellbore\\" + name, low_memory=False) for name in names], ignore_index=True).drop_duplicates("id") for month, names in files.items()}
    frames = {}
    for month, frame in raw.items():
        frame = frame.loc[frame["currentAccount"].isin(["3-1-robot", "3-1-robot2"])].copy()
        frame["month"] = month
        frame["datetime"] = pd.to_datetime(frame["createTime"], errors="coerce")
        frame["date"] = frame["datetime"].dt.normalize()
        frame["score_num"] = pd.to_numeric(frame["score"], errors="coerce")
        frame["rated"] = frame["score_num"].between(1, 5)
        frame["good"] = frame["score_num"].isin([3, 4, 5])
        frame["from_alika"] = frame["changeFlag"].astype("string").str.contains("Alika", case=False, na=False)
        last = frame.groupby("date")["datetime"].max()
        frame = frame.loc[frame["date"].isin(last.loc[last.dt.hour.ge(23)].index)]
        frames[month] = frame
    monthly = pd.DataFrame({month: {"volume": len(frame), "days": frame["date"].nunique(), "respondents": int(frame["rated"].sum()), "good": 100 * frame.loc[frame["rated"], "good"].mean(), "bad": 100 * (1 - frame.loc[frame["rated"], "good"].mean()), "csat": frame.loc[frame["rated"], "score_num"].mean()} for month, frame in frames.items()}).T
    colors = ["#3976a8", "#3976a8", "#3976a8", "#d77932"]
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
    axes[0].bar(monthly.index, monthly["volume"], color=colors)
    axes[0].set_title("Total KULA conversations")
    axes[0].set_ylabel("Conversations")
    axes[1].bar(monthly.index, monthly["volume"] / monthly["days"], color=colors)
    axes[1].set_title("Average per complete day")
    axes[1].set_ylabel("Conversations/day")
    fig.suptitle("KULA Complaint Volume, May-August 2026", fontweight="bold")
    fig.tight_layout()
    volume_path = FIGURE_DIR + "\\figure1_volume.png"
    fig.savefig(volume_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))
    axes[0].bar(monthly.index, monthly["good"], color="#4c9f70", label="Good (3-5)")
    axes[0].bar(monthly.index, monthly["bad"], bottom=monthly["good"], color="#c95b58", label="Bad (1-2)")
    axes[0].set_title("Survey quality")
    axes[0].set_ylabel("Percent of valid respondents")
    axes[0].legend(frameon=False)
    distribution = pd.DataFrame({month: frame.loc[frame["rated"], "score_num"].value_counts(normalize=True).mul(100) for month, frame in frames.items()}).fillna(0)
    distribution.plot.bar(ax=axes[1], color=colors)
    axes[1].set_title("Rating score distribution")
    axes[1].set_xlabel("Score")
    axes[1].set_ylabel("Percent of valid respondents")
    axes[1].legend(frameon=False)
    fig.suptitle("KULA Survey Results, May-August 2026", fontweight="bold")
    fig.tight_layout()
    quality_path = FIGURE_DIR + "\\figure2_survey_quality.png"
    fig.savefig(quality_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    all_data = pd.concat(frames.values(), ignore_index=True)
    monthly_forwarding = all_data.groupby("month").agg(
        alika_forwarded=("from_alika", "sum"),
        kula_conversations=("id", "size"),
    ).reindex(["May", "June", "July", "August"])
    monthly_forwarding["forwarded_pct"] = 100 * monthly_forwarding["alika_forwarded"] / monthly_forwarding["kula_conversations"]
    fig, axis = plt.subplots(figsize=(12, 4.2))
    bars = axis.bar(monthly_forwarding.index, monthly_forwarding["forwarded_pct"], color=["#3976a8", "#3976a8", "#3976a8", "#d77932"])
    axis.bar_label(bars, labels=[f"{count:,} ({pct:.3f}%)" for count, pct in zip(monthly_forwarding["alika_forwarded"], monthly_forwarding["forwarded_pct"])], padding=3, fontsize=8)
    axis.set_title("Monthly KULA Conversations Marked as Forwarded from ALIKA", fontweight="bold")
    axis.set_ylabel("ALIKA-forwarded conversations (% of KULA volume)")
    axis.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    forwarding_path = FIGURE_DIR + "\\figure3_monthly_alika_forwarding.png"
    fig.savefig(forwarding_path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return volume_path, quality_path, forwarding_path


FIGURES = build_figures()


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[index], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph()
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(document, text, italic=False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.italic = italic
    return paragraph


def add_figure(document, path, caption, explanation):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(path, width=Inches(6.7))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.bold = True
    caption_run.font.size = Pt(9)
    add_body(document, explanation)


document = Document()
section = document.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = document.styles
styles["Normal"].font.name = "Georgia"
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.name = "Georgia"
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 2"].font.name = "Georgia"
styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run("Observed Association Between ALIKA Routing and KULA Service Outcomes")
run.bold = True
run.font.name = "Georgia"
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("A four-month retrospective analysis of KULA conversations, survey quality, and routing markers").italic = True
byline = document.add_paragraph()
byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
byline.add_run("Analysis period: 1 May-31 August 2026 | Report date: 3 September 2026")

add_heading(document, "Abstract", 1)
add_body(document, "Background: This study assessed whether KULA service outcomes changed during the period in which conversations were observed to route from ALIKA. Methods: KULA records from May through August 2026 were deduplicated by conversation identifier, restricted to the account labels 3-1-robot and 3-1-robot2, and compared across complete export dates. Valid survey scores were 1-5; good surveys were scores 3-5. The primary cutoff was the first explicit ALIKA routing marker, 31 July 2026. Welch's t-test compared mean CSAT and a Pearson chi-square test compared good/bad survey composition. Results: The observed-routing period had 190,417 conversations across 29 complete days versus 643,117 across 91 days before routing. Mean CSAT increased from 3.250 to 3.361 (+0.111, p=0.038), while good-survey rate increased from 58.70% to 60.93% (+2.23 percentage points, p=0.113). Survey response rate decreased from 0.992% to 0.790%. Monthly ALIKA-marked forwarding was 0 conversations in May, 0 in June, 1 in July, and 772 in August; August's 772 marked conversations represented 0.423% of complete-date KULA volume. Conclusion: KULA outcomes improved in the observed routing period, but the evidence does not establish that ALIKA caused the improvement. Sparse and delayed routing markers, falling response rates, temporal confounding, and low power for the good-survey endpoint limit causal interpretation.")
add_body(document, "Keywords: customer satisfaction; CSAT; chatbot routing; before-after study; service operations; observational analysis", italic=True)

add_heading(document, "1. Introduction", 1)
add_body(document, "Operational teams often evaluate a conversational-system rollout by comparing customer feedback before and after routing changes. Such comparisons are useful for monitoring, but they can overstate causal effects when rollout timing is uncertain, exposure is incompletely recorded, or customer composition changes over time. This analysis examines the relationship between ALIKA routing and KULA outcomes using the available export data.")
add_body(document, "The primary analysis uses the earliest complete-date KULA record whose changeFlag explicitly indicates transfer from ALIKA, because that is an observable data event. The study question is therefore deliberately narrower than a causal launch evaluation: did KULA volume and survey outcomes differ after the first observed ALIKA routing marker?")

add_heading(document, "2. Methods", 1)
add_heading(document, "2.1 Design and data preparation", 2)
add_body(document, "This was a retrospective, uncontrolled before-after study. Twenty-six May-July exports and eight August exports were concatenated within month and deduplicated by id. KULA traffic was defined as currentAccount in {3-1-robot, 3-1-robot2}. Dates whose latest available record did not reach hour 23 were excluded from normalized daily and phase comparisons; 18, 20, and 31 August were incomplete and excluded. This yielded 833,534 complete-date KULA conversations across the four months.")
add_heading(document, "2.2 Measures", 2)
add_body(document, "Conversation volume was the number of unique deduplicated conversations. Survey respondents were records with numeric scores from 1 through 5. CSAT was calculated by adding all valid survey scores and dividing by the number of valid respondents. For example, five valid scores of 5, 4, 3, 2, and 1 sum to 15; dividing by five gives a CSAT of 3.0. This is the average rating among people who answered, not a percentage and not a score assigned to conversations without a survey response. Good-survey rate was the proportion of valid respondents scoring 3-5, and bad-survey rate was the proportion scoring 1-2. Survey response rate was valid respondents divided by conversations. The routing marker was a case-insensitive search for Alika in changeFlag. Monthly ALIKA forwarding was the number of complete-date KULA conversations with that marker and its percentage of all complete-date KULA conversations in the month.")
add_heading(document, "2.3 Statistical analysis", 2)
add_body(document, "The primary comparison used Welch's independent-samples t-test for CSAT and an uncorrected Pearson chi-square test for the 2-by-2 good/bad survey table, with alpha=0.05. Monthly differences were assessed with one-way ANOVA, Kruskal-Wallis, and a monthly good/bad chi-square test. A sensitivity analysis used 28 June as the cutoff. A post hoc two-proportion power calculation estimated the additional routing-period respondents needed to reach approximately 80% power for the observed good-survey difference, holding the pre-period sample fixed. Analyses were performed in Python with pandas, NumPy, SciPy, and Matplotlib/Seaborn.")

add_heading(document, "3. Results", 1)
add_heading(document, "3.1 Data coverage and launch validation", 2)
add_table(document, ["Month", "Conversations", "Complete days", "Respondents", "CSAT", "Good %", "Response %"], [
    ["May", "240,894", "31", "2,551", "3.148", "56.10", "1.059"],
    ["June", "198,542", "30", "2,103", "3.238", "58.11", "1.059"],
    ["July", "211,595", "31", "1,780", "3.417", "63.32", "0.841"],
    ["August", "182,503", "28", "1,451", "3.357", "60.79", "0.795"],
], [0.8, 1.2, 1.0, 1.0, 0.7, 0.8, 0.9])
add_body(document, "The first ALIKA-marked KULA conversation occurred at 09:54:24 on 31 July 2026. Monthly marked forwarding was 0 of 240,894 KULA conversations in May, 0 of 198,542 in June, 1 of 211,595 in July, and 772 of 182,503 in August. Thus, the marker was concentrated in August and represented 0.423% of August's complete-date KULA volume. This makes the marker a sparse exposure indicator rather than a complete treatment assignment.")

add_heading(document, "3.2 Monthly ALIKA forwarding", 2)
add_table(document, ["Month", "ALIKA-forwarded", "KULA conversations", "Forwarded %"], [
    ["May", "0", "240,894", "0.000%"],
    ["June", "0", "198,542", "0.000%"],
    ["July", "1", "211,595", "0.000%"],
    ["August", "772", "182,503", "0.423%"],
], [1.0, 1.4, 1.5, 1.0])
add_body(document, "This table counts only conversations whose changeFlag explicitly contains ALIKA. The percentage answers how much of each month's KULA traffic was visibly marked as forwarded from ALIKA. It should not be interpreted as the total share of ALIKA-influenced conversations, because an unmarked conversation may still have had exposure that was not recorded in this field.")

add_heading(document, "3.2 Primary before-versus-observed-routing comparison", 2)
add_table(document, ["Outcome", "Before (through 30 Jul)", "Routing period (from 31 Jul)", "Difference", "p-value"], [
    ["Conversations/day", "7,067.2", "6,566.1", "-501.1 (-7.1%)", "Descriptive"],
    ["Valid respondents", "6,380", "1,505", "-4,875", "Descriptive"],
    ["Survey response rate", "0.992%", "0.790%", "-0.202 pp", "Descriptive"],
    ["CSAT mean", "3.250", "3.361", "+0.111", "0.038"],
    ["Good survey rate", "58.70%", "60.93%", "+2.23 pp", "0.113"],
], [1.5, 1.3, 1.5, 1.4, 0.7])
add_body(document, "The CSAT difference was statistically significant at the nominal 0.05 level. The good-survey proportion was directionally higher but did not meet that threshold. Because the analysis tested multiple related outcomes and did not adjust for multiplicity, the CSAT p-value should be interpreted as exploratory rather than confirmatory.")

add_heading(document, "3.3 Monthly findings", 2)
add_body(document, "Outcomes varied across months: the monthly CSAT means were 3.148, 3.238, 3.417, and 3.357 from May through August. The monthly omnibus tests were significant for CSAT by ANOVA (F=8.555, p<0.001), score distributions by Kruskal-Wallis (H=24.369, p<0.001), and good/bad composition by chi-square (chi-square=25.172, p<0.001). These findings show time variation, but they do not identify ALIKA as the cause of that variation.")
add_body(document, "The first observed ALIKA routing marker occurred late in the study period, so the before-versus-routing comparison has unequal time windows. The monthly pattern also shows that survey outcomes changed over time independently of the sparse forwarding marker. These results should therefore be treated as monitoring evidence rather than proof of an ALIKA effect.")

add_heading(document, "3.4 Precision and power", 2)
add_body(document, "For the observed +2.23 percentage-point good-survey difference, the available sample provided approximately 35.3% power under the notebook's two-proportion approximation. Approximately 9,316 routing-period respondents would be required for 80% power with the pre-period sample fixed; 7,811 additional respondents were estimated, equivalent to about 151 additional complete days at the observed rate of 51.9 respondents per day. Thus, one or two additional 30-day months would not be expected to reach 80% power under the same assumptions; the notebook estimates approximately six additional 30-day months.")

add_heading(document, "3.4 Visual analysis", 2)
add_figure(document, FIGURES[0], "Figure 1. KULA complaint volume by month and average volume per complete day.", "Volume declined from 240,894 conversations in May to 182,503 complete-date conversations in August. The daily view shows that this is not only a calendar-length effect: average volume fell from 7,771 conversations per day in May to 6,518 in August. The observed-routing period averaged 6,566 conversations per day, 7.1% below the preceding period. This operational change is a potential confounder when interpreting survey outcomes.")
add_figure(document, FIGURES[1], "Figure 2. Monthly survey quality and valid-rating score distribution.", "The share of good surveys rose from 56.10% in May to 63.32% in July, then eased to 60.79% in August. The score distribution shows that the improvement is largely reflected in a lower share of score-1 responses and a higher share of score-5 responses in July and August. Because response rates also decline over the same period, the chart describes respondents rather than all conversations and may be affected by non-response selection.")
add_figure(document, FIGURES[2], "Figure 3. Monthly KULA conversations explicitly marked as forwarded from ALIKA.", "The monthly chart shows no marked forwarding in May or June, one marked conversation in July, and 772 in August. Even in August, the marked conversations represented only 0.423% of complete-date KULA volume. This is useful operational evidence about recorded forwarding, but it is too sparse to represent all ALIKA exposure.")

add_heading(document, "4. Discussion", 1)
add_body(document, "The data show a modest improvement in average KULA survey score after 31 July, alongside a non-significant increase in the good-survey rate. The pattern is compatible with improved customer experience, but the design cannot separate an ALIKA effect from calendar trends, changes in traffic mix, operational changes, differences in survey participation, or regression to the mean. The concurrent 7.1% decline in daily conversation volume and 0.202 percentage-point decline in survey response rate are especially important: the observed respondents may not represent the same customer population before and after the cutoff.")
add_body(document, "The routing marker is not suitable for a direct exposed-versus-unexposed causal comparison in this extract. Only 773 conversations were marked from ALIKA across the observed-routing period, and the marked subgroup had just two valid survey respondents; their CSAT and survey proportions are therefore unstable. The first non-null autoTransferType also appeared earlier, on 14 May, and only two such records were present, so it does not provide a reliable rollout boundary.")
add_heading(document, "4.1 Strengths", 2)
add_body(document, "Strengths include explicit deduplication, inclusion of both KULA account labels, exclusion of incomplete export dates, separation of survey quality from complaint volume, use of a routing-marker validation step, and a monthly count of recorded ALIKA forwarding.")
add_heading(document, "4.2 Limitations", 2)
add_body(document, "Limitations include the uncontrolled before-after design; uncertain and sparse exposure measurement; unequal phase lengths; possible changes in customer mix and operations; low and declining survey response; repeated observations that may not be independent at the user level; no confidence intervals or multiplicity correction in the notebook; and the stale-kernel failure of the final machine-readable conclusion cell. The numerical tables used here were generated successfully from the preceding analytical cells, but the notebook should be rerun from the top for a clean reproducibility record.")

add_heading(document, "5. Conclusion", 1)
add_body(document, "KULA CSAT was higher during the period beginning with the first observed ALIKA routing marker, while the good-survey rate showed a smaller, statistically non-significant increase. The evidence supports monitoring the result as a promising operational signal, not claiming a causal ALIKA impact. A stronger evaluation should record ALIKA exposure at conversation level, use a validated rollout timestamp, preserve a contemporaneous comparison group, and continue collecting surveys until the primary endpoint has adequate precision.")

add_heading(document, "References", 1)
references = [
    "1. von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet. 2007;370:1453-1457.",
    "2. Welch BL. The generalization of Student's problem when several different population variances are involved. Biometrika. 1947;34(1-2):28-35.",
    "3. Pearson K. On the criterion that a given system of deviations from the probable in the case of a correlated system of variables is such that it can be reasonably supposed to have arisen from random sampling. Philosophical Magazine. 1900;50(302):157-175.",
    "4. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. Lawrence Erlbaum Associates; 1988.",
]
for reference in references:
    add_body(document, reference)

add_heading(document, "Appendix A. Reproducibility note", 1)
add_body(document, "Source notebook: notebooks/alika_effect_on_kula.ipynb. Source data: dataset_wellbore monthly exports referenced in that notebook. The report reflects the notebook's definitions and calculations as executed in the current environment on 3 September 2026. Values are rounded for presentation; p-values shown as <0.001 correspond to values below 0.001 in the executed output.")

for paragraph in document.paragraphs:
    paragraph.paragraph_format.widow_control = True

document.save(OUTPUT)
print(OUTPUT)